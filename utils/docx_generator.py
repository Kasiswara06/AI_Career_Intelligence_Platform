import os
import logging
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

def generate_resume_docx(resume_dict: dict, output_filepath: str) -> str:
    """
    Generates a professional DOCX resume file using python-docx.
    """
    doc = docx.Document()

    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    # Name Header
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_name = p_name.add_run(resume_dict.get("full_name", "Candidate Name").upper())
    run_name.font.size = Pt(20)
    run_name.font.bold = True
    run_name.font.color.rgb = RGBColor(15, 23, 42)

    # Contact Line
    contact_parts = [p for p in [resume_dict.get('email'), resume_dict.get('phone'), resume_dict.get('location')] if p]
    if contact_parts:
        p_contact = doc.add_paragraph()
        p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_contact = p_contact.add_run(" | ".join(contact_parts))
        run_contact.font.size = Pt(9.5)
        run_contact.font.color.rgb = RGBColor(71, 85, 105)

    # Social Links
    link_parts = []
    if resume_dict.get("linkedin"): link_parts.append(f"LinkedIn: {resume_dict['linkedin']}")
    if resume_dict.get("github"): link_parts.append(f"GitHub: {resume_dict['github']}")
    if link_parts:
        p_links = doc.add_paragraph()
        p_links.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_links = p_links.add_run(" | ".join(link_parts))
        run_links.font.size = Pt(9)
        run_links.font.color.rgb = RGBColor(71, 85, 105)

    # Helper for adding sections
    def add_section(title, content):
        if not content or not content.strip():
            return
        p_sec = doc.add_paragraph()
        p_sec.paragraph_format.space_before = Pt(12)
        p_sec.paragraph_format.space_after = Pt(4)
        run_sec = p_sec.add_run(title.upper())
        run_sec.font.size = Pt(11)
        run_sec.font.bold = True
        run_sec.font.color.rgb = RGBColor(37, 99, 235)

        for line in content.split("\n"):
            if line.strip():
                p_item = doc.add_paragraph()
                p_item.paragraph_format.space_after = Pt(2)
                run_item = p_item.add_run(line.strip())
                run_item.font.size = Pt(10)
                run_item.font.color.rgb = RGBColor(51, 65, 85)

    add_section("Professional Summary", resume_dict.get("summary"))
    add_section("Technical Skills", resume_dict.get("skills"))
    add_section("Education", resume_dict.get("education"))
    add_section("Professional Experience", resume_dict.get("experience"))
    add_section("Key Projects", resume_dict.get("projects"))
    add_section("Certifications", resume_dict.get("certifications"))
    add_section("Achievements", resume_dict.get("achievements"))

    os.makedirs(os.path.dirname(output_filepath), exist_ok=True)
    doc.save(output_filepath)
    return output_filepath
